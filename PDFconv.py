import os
import fitz
from docx import Document
from pdf2image import convert_from_path
from concurrent.futures import ThreadPoolExecutor

def extract_text_from_page(page):
    return page.get_text()

def convert_page_to_images(pdf_path, page_num, image_folder):
    images = convert_from_path(pdf_path, first_page=page_num, last_page=page_num)
    image_paths = []
    for image_num, image in enumerate(images):
        image_path = os.path.join(image_folder, f"page_{page_num}_image_{image_num}.png")
        image.save(image_path)
        image_paths.append(image_path)
    return image_paths

def pdf_to_word(pdf_path, output_path, max_image_width=400):  # Adjust max_image_width as needed
    # Create a folder for images
    image_folder = os.path.join(os.path.dirname(output_path), "images")
    os.makedirs(image_folder, exist_ok=True)

    doc = fitz.open(pdf_path)

    # Create a new Word document
    word_doc = Document()

    with ThreadPoolExecutor() as executor:
        text_futures = [executor.submit(extract_text_from_page, page) for page in doc]
        image_futures = [executor.submit(convert_page_to_images, pdf_path, i, image_folder) for i in range(len(doc))]

        for future in text_futures:
            text = future.result()
            word_doc.add_paragraph(text)

        for future in image_futures:
            image_paths = future.result()
            for image_path in image_paths:
                # Add the image to the Word document with specified width
                word_doc.add_picture(image_path, width=max_image_width)

    # Save the Word document
    word_doc.save(output_path)

# Specify the paths
pdf_path = "The Linux Command Line.pdf"
output_path = "output.docx"

# Convert PDF to Word with text and images (setting max_image_width to 400)
pdf_to_word(pdf_path, output_path, max_image_width=400)
